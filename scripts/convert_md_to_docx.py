from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


def apply_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size_pt in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size_pt)


def format_paragraph(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    apply_document_styles(doc)

    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            p = doc.add_paragraph("")
            format_paragraph(p)
            continue

        if stripped == "---":
            doc.add_page_break()
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            format_paragraph(p)
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            format_paragraph(p)
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            format_paragraph(p)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            format_paragraph(p)
            continue

        if stripped.startswith("1. "):
            p = doc.add_paragraph(stripped[3:].strip(), style="List Number")
            format_paragraph(p)
            continue

        p = doc.add_paragraph(line)
        format_paragraph(p)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> None:
    workspace = Path(__file__).resolve().parents[1]
    md_path = workspace / "docs" / "THIEL_FELLOWSHIP_APPLICATION_DRAFT.md"
    docx_path = workspace / "docs" / "THIEL_FELLOWSHIP_APPLICATION_DRAFT.docx"

    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    convert_markdown_to_docx(md_path, docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    main()
